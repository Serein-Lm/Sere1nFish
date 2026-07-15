"""
通用 Word 文档生成服务。

将 Markdown 风格文本或结构化段落渲染为内存中的 .docx 字节。
纯文档生成不涉及存储；上传和元信息登记由统一 StorageService 与 DAO 完成。

设计为「通用工具」：不限定内容主体，AI 中枢 agent 可直接调用生成任意 Word 交给前端下载。
python-docx 采用惰性导入，未安装时抛出清晰错误由上层捕获。
"""
from __future__ import annotations

import re
import uuid
from io import BytesIO
from typing import Any


def _safe_filename(title: str) -> str:
    """从标题生成安全的文件名主体。"""
    base = re.sub(r"[^\w\u4e00-\u9fa5\-]+", "_", (title or "document").strip())
    base = base.strip("_") or "document"
    return base[:60]


def _apply_markdown_line(doc: Any, line: str) -> None:
    """将单行 Markdown 文本映射为 docx 段落/标题/列表项。"""
    stripped = line.rstrip()
    if not stripped.strip():
        doc.add_paragraph("")
        return

    heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
    if heading:
        level = min(len(heading.group(1)), 4)
        doc.add_heading(heading.group(2).strip(), level=level)
        return

    if re.match(r"^\s*[-*+]\s+", stripped):
        text = re.sub(r"^\s*[-*+]\s+", "", stripped)
        doc.add_paragraph(_strip_inline(text), style="List Bullet")
        return

    ordered = re.match(r"^\s*\d+[.)]\s+(.*)$", stripped)
    if ordered:
        doc.add_paragraph(_strip_inline(ordered.group(1)), style="List Number")
        return

    if re.match(r"^\s*>\s+", stripped):
        text = re.sub(r"^\s*>\s+", "", stripped)
        doc.add_paragraph(_strip_inline(text), style="Intense Quote")
        return

    doc.add_paragraph(_strip_inline(stripped))


def _strip_inline(text: str) -> str:
    """去除行内 Markdown 标记（加粗/斜体/行内代码），保留纯文本。"""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def generate_docx(
    *,
    title: str,
    content: str = "",
    sections: list[dict[str, str]] | None = None,
    owner: str = "",
) -> dict[str, Any]:
    """
    生成 Word 文档字节。

    参数:
        title: 文档标题（作为一级标题与文件名基）。
        content: Markdown 风格正文（与 sections 二选一或叠加）。
        sections: 结构化段落列表，每项 {"heading": str, "body": str}。
        owner: 归属用户名（写入元信息用）。

    返回:
        dict: artifact_id / data / filename / size / title / kind
    """
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError(
            "未安装 python-docx，无法生成 Word 文档。请在服务端执行 pip install python-docx。"
        ) from exc

    doc = Document()
    doc.add_heading(title or "未命名文档", level=0)

    if content and content.strip():
        for line in content.splitlines():
            _apply_markdown_line(doc, line)

    for section in sections or []:
        heading = (section.get("heading") or "").strip()
        body = section.get("body") or ""
        if heading:
            doc.add_heading(heading, level=1)
        for line in body.splitlines():
            _apply_markdown_line(doc, line)

    artifact_id = "art_" + uuid.uuid4().hex[:20]
    filename = f"{_safe_filename(title)}.docx"
    buffer = BytesIO()
    doc.save(buffer)
    data = buffer.getvalue()

    return {
        "artifact_id": artifact_id,
        "data": data,
        "filename": filename,
        "size": len(data),
        "title": title or "未命名文档",
        "kind": "word",
    }
