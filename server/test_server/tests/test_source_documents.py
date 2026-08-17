"""来源文档版本、Target 聚类与联系方式证据的纯逻辑回归测试。"""

import pytest


def test_docx_embedded_raster_images_use_bounded_ocr(monkeypatch):
    import io
    import zipfile
    from types import SimpleNamespace

    from api.services.source_documents import resources

    payload = io.BytesIO()
    with zipfile.ZipFile(payload, "w") as archive:
        archive.writestr("word/media/image1.jpeg", b"jpeg-payload")

    monkeypatch.setattr(resources.shutil, "which", lambda name: f"/usr/bin/{name}")
    monkeypatch.setattr(
        resources.subprocess,
        "run",
        lambda *_args, **_kwargs: SimpleNamespace(
            returncode=0,
            stdout="联系人 010-12345678".encode(),
            stderr=b"",
        ),
    )

    text, error = resources._ocr_docx_images(payload.getvalue(), 10_000)

    assert "010-12345678" in text
    assert error == ""


def test_docx_unsupported_vector_images_remain_explicitly_partial(monkeypatch):
    import io
    import zipfile

    from api.services.source_documents import resources

    payload = io.BytesIO()
    with zipfile.ZipFile(payload, "w") as archive:
        archive.writestr("word/media/image1.emf", b"emf-payload")

    monkeypatch.setattr(resources.shutil, "which", lambda name: f"/usr/bin/{name}")

    text, error = resources._ocr_docx_images(payload.getvalue(), 10_000)

    assert text == ""
    assert "emf" in error


def test_docx_vector_scan_falls_back_to_office_pdf(monkeypatch):
    from api.services.source_documents import resources

    monkeypatch.setattr(
        resources,
        "_ocr_docx_images",
        lambda _data, _limit: ("", "unsupported emf"),
    )
    monkeypatch.setattr(
        resources,
        "_extract_office_via_pdf",
        lambda _data, *, suffix, limit: (
            f"converted {suffix} {limit}",
            "",
        ),
    )

    import io
    import zipfile

    payload = io.BytesIO()
    with zipfile.ZipFile(payload, "w") as archive:
        archive.writestr(
            "[Content_Types].xml",
            """<?xml version="1.0" encoding="UTF-8"?>
            <Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
              <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
              <Default Extension="xml" ContentType="application/xml"/>
              <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
            </Types>""",
        )
        archive.writestr(
            "_rels/.rels",
            """<?xml version="1.0" encoding="UTF-8"?>
            <Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
              <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
            </Relationships>""",
        )
        archive.writestr(
            "word/document.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
            <w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body/></w:document>""",
        )

    text, error = resources._extract_docx_text(payload.getvalue(), 12_345)

    assert text == "converted .docx 12345"
    assert error == ""


def test_image_only_legacy_xls_falls_back_to_office_pdf(monkeypatch):
    import sys
    from types import SimpleNamespace

    from api.services.source_documents import resources

    class _Workbook:
        def sheets(self):
            return [
                SimpleNamespace(
                    name="Sheet1",
                    nrows=0,
                    ncols=0,
                    cell_value=lambda *_args: "",
                )
            ]

        def release_resources(self):
            return None

    monkeypatch.setitem(
        sys.modules,
        "xlrd",
        SimpleNamespace(open_workbook=lambda **_kwargs: _Workbook()),
    )
    monkeypatch.setattr(
        resources,
        "_extract_office_via_pdf",
        lambda _data, *, suffix, limit: (
            f"OCR 后正文 {suffix} {limit}",
            "",
        ),
    )

    text, error, text_format = resources.extract_attachment_text(
        b"legacy-xls",
        filename="扫描表格.xls",
        content_type="application/vnd.ms-excel",
        max_chars=12_345,
    )

    assert text == "工作表: Sheet1 OCR 后正文 .xls 12345"
    assert error == ""
    assert text_format == "xls"


def test_docx_short_caption_does_not_hide_embedded_scan(monkeypatch):
    import io

    from docx import Document

    from api.services.source_documents import resources

    document = Document()
    document.add_paragraph("附件一")
    document.add_picture(io.BytesIO(_one_pixel_png()))
    payload = io.BytesIO()
    document.save(payload)

    monkeypatch.setattr(
        resources,
        "_ocr_docx_images",
        lambda _data, _limit: ("联系人 010-12345678", ""),
    )

    text, error = resources._extract_docx_text(payload.getvalue(), 10_000)

    assert "附件一" in text
    assert "010-12345678" in text
    assert error == ""


def _one_pixel_png() -> bytes:
    import base64

    return base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUB"
        "AScY42YAAAAASUVORK5CYII="
    )


def _capture(*, raw_html: bytes, rendered_html: bytes):
    from api.services.source_documents.contracts import (
        CapturedDocument,
        CapturedImage,
    )

    return CapturedDocument(
        source_type="wechat_article",
        canonical_url="https://mp.weixin.qq.com/s/article-id",
        requested_url="https://mp.weixin.qq.com/s/article-id?scene=1",
        title="测试文章",
        account="测试公众号",
        publish_time="2026-07-16",
        text="联系人张三，手机 13800138000。",
        raw_html=raw_html,
        rendered_html=rendered_html,
        images=[
            CapturedImage(
                index=0,
                source_url="https://mmbiz.qpic.cn/image-1",
                data=b"image-one",
                content_type="image/jpeg",
            )
        ],
    )


def test_text_analysis_uses_collection_model_and_reports_actual_model(monkeypatch):
    import asyncio
    from types import SimpleNamespace

    from api.services.source_documents import analysis

    outputs = iter(
        [
            {
                "summary": "目标文章摘要",
                "article_scope": "target_focused",
                "target_contact_values": [],
                "subject_match": 95,
                "relevance_score": 90,
                "score_reason": "主体明确",
            },
            {
                "decision": "accept",
                "article_scope": "target_focused",
                "subject_match": 95,
                "relevance_score": 90,
                "summary": "目标文章摘要",
                "target_contact_values": [],
                "reason": "主体明确",
            },
        ]
    )
    calls: list[dict] = []

    class _StructuredLLM:
        async def ainvoke(self, _messages):
            return next(outputs)

    class _FakeLLM:
        model_name = "deepseek-v4-flash-0731"

        def with_structured_output(self, _schema):
            return _StructuredLLM()

    async def _runtime_config():
        return SimpleNamespace(
            runtime=SimpleNamespace(
                models=SimpleNamespace(
                    default="qwen3.8-max",
                    collection_model="deepseek-v4-flash-0731",
                )
            )
        )

    def _create_llm(_app_config, **kwargs):
        calls.append(kwargs)
        return _FakeLLM()

    monkeypatch.setattr(analysis, "get_runtime_app_config", _runtime_config)
    monkeypatch.setattr(analysis, "create_llm", _create_llm)
    monkeypatch.setattr(analysis, "load_prompt", lambda _slug: "prompt")
    capture = _capture(raw_html=b"raw", rendered_html=b"rendered")

    async def run():
        structured = await analysis.analyze_article_fields(
            capture,
            fields=[],
            target_name="测试单位",
            keyword="测试单位 联系方式",
        )
        review = await analysis.review_article_relevance(
            capture,
            draft_analysis=structured,
            target_name="测试单位",
            keyword="测试单位 联系方式",
            required_subject_match=70,
        )
        return structured, review

    structured, review = asyncio.run(run())

    assert structured["analysis_model"] == "deepseek-v4-flash-0731"
    assert review["review_model"] == "deepseek-v4-flash-0731"
    assert [call["workload"] for call in calls] == ["collection", "collection"]


def test_wechat_canonical_url_discards_tracking_query_and_fragment():
    from api.services.source_documents.urls import canonicalize_source_url

    canonical = canonicalize_source_url(
        "http://MP.WEIXIN.QQ.COM/s/article-id?scene=1&from=timeline#wechat_redirect"
    )
    assert canonical == "https://mp.weixin.qq.com/s/article-id"


def test_stable_content_hash_ignores_dynamic_page_shell():
    from api.dao.source_documents import (
        document_id_for_url,
        version_id_for_content,
    )
    from api.services.source_documents.analysis import stable_content_hash

    first = _capture(raw_html=b"token=one", rendered_html=b"session=one")
    second = _capture(raw_html=b"token=two", rendered_html=b"session=two")

    assert stable_content_hash(first) == stable_content_hash(second)

    document_id = document_id_for_url(first.canonical_url)
    assert version_id_for_content(
        document_id, stable_content_hash(first)
    ) == version_id_for_content(document_id, stable_content_hash(second))


def test_stable_hash_uses_declared_images_when_a_download_is_partial():
    from api.services.source_documents.analysis import stable_content_hash

    complete = _capture(raw_html=b"raw-one", rendered_html=b"dom-one")
    complete.metadata["image_urls"] = ["https://mmbiz.qpic.cn/image-1"]
    partial = _capture(raw_html=b"raw-two", rendered_html=b"dom-two")
    partial.metadata["image_urls"] = ["https://mmbiz.qpic.cn/image-1"]
    partial.images = []

    assert stable_content_hash(complete) == stable_content_hash(partial)


def test_stable_hash_uses_declared_attachments_when_download_is_partial():
    from api.services.source_documents.analysis import stable_content_hash
    from api.services.source_documents.contracts import CapturedAttachment

    complete = _capture(raw_html=b"raw-one", rendered_html=b"dom-one")
    complete.metadata.update(
        {
            "article_text": complete.text,
            "attachment_urls": ["https://example.com/files/contact.xlsx"],
        }
    )
    complete.text += "\n附件提取出的动态文本"
    complete.attachments = [
        CapturedAttachment(
            index=0,
            source_url="https://example.com/files/contact.xlsx",
            filename="contact.xlsx",
            data=b"xlsx",
            content_type="application/octet-stream",
            sha256="downloaded",
            extracted_text="附件提取出的动态文本",
        )
    ]
    partial = _capture(raw_html=b"raw-two", rendered_html=b"dom-two")
    partial.metadata.update(
        {
            "article_text": partial.text,
            "attachment_urls": ["https://example.com/files/contact.xlsx"],
        }
    )
    partial.attachments = []

    assert stable_content_hash(complete) == stable_content_hash(partial)


def test_cached_refresh_preserves_complete_archived_attachment_text():
    from dataclasses import replace

    from api.services.source_documents import service

    capture = replace(
        _capture(raw_html=b"raw", rendered_html=b"dom"),
        text="公告正文",
        attachments=[],
        metadata={
            "attachment_urls": ["https://example.gov.cn/contact.pdf"],
            "attachment_download_errors": ["timeout"],
        },
    )
    archived_text = "公告正文\n附件联系人：archive@example.gov.cn"
    preserved = service._preserve_cached_evidence_text(
        {
            "content": {"text": archived_text},
            "attachments": [
                {
                    "source_url": "https://example.gov.cn/contact.pdf",
                    "text_length": len(archived_text),
                }
            ],
        },
        capture,
    )

    assert preserved.text == archived_text


def test_generic_source_url_discards_tracking_parameters():
    from api.services.source_documents.urls import canonicalize_source_url

    assert canonicalize_source_url(
        "https://Example.com/notice?id=7&utm_source=feed&from=share#top"
    ) == "https://example.com/notice?id=7"


def test_source_document_identity_prevents_field_change_duplicates():
    from api.dao.mobile_collect import stable_record_id

    first = stable_record_id(
        "task-1",
        {"title": "旧标题"},
        ["title"],
        source_document_id="doc-1",
    )
    second = stable_record_id(
        "task-1",
        {"title": "新标题", "summary": "新结构化结果"},
        ["title"],
        source_document_id="doc-1",
    )
    assert first == second


def test_target_identity_normalizes_common_name_separators():
    from api.dao.targets import normalize_target_name, target_id_for_name

    assert normalize_target_name(" 天津-滨海 国际机场（集团） ") == "天津滨海国际机场集团"
    assert target_id_for_name("天津-滨海 国际机场") == target_id_for_name(
        "天津滨海国际机场"
    )


def test_contact_extraction_preserves_local_evidence_context():
    from core.mobile.collect.contacts import extract_contacts

    text = (
        "项目报名要求如下。\n"
        "商务联系人：张三，手机 13800138000，负责投标资料接收。\n"
        "其他事项请查看原文。"
    )
    contacts = extract_contacts(text)

    phone = next(item for item in contacts if item["channel"] == "phone")
    assert phone["value"] == "13800138000"
    assert "张三" in phone["context"]
    assert "投标资料接收" in phone["context"]
    assert phone["contexts"] == [phone["context"]]


def test_contact_extraction_reconstructs_tabular_row_context():
    from core.mobile.collect.contacts import extract_contacts

    text = (
        "院系代码\n院系简称\n联系人\n联系方式\n邮箱\n\n"
        "001\n中心总部\n田老师\n010-58900576\ntianyy@chinacdc.cn\n\n"
        "002\n传染病所\n崔老师\n010-58900713\ncuiyao@icdc.cn"
    )
    contacts = extract_contacts(text)

    email = next(item for item in contacts if item["value"] == "tianyy@chinacdc.cn")
    phone = next(item for item in contacts if item["value"] == "010-58900576")
    assert "中心总部" in email["context"]
    assert "田老师" in email["context"]
    assert "中心总部" in phone["context"]
    assert "田老师" in phone["context"]


def test_contact_extraction_normalizes_full_width_email_and_service_phone():
    from core.mobile.collect.contacts import extract_contacts

    contacts = extract_contacts(
        "监督邮箱：lwzx_jw＠126.com；服务热线：400 123 4567。"
    )

    assert [(item["channel"], item["value"]) for item in contacts] == [
        ("email", "lwzx_jw@126.com"),
        ("telephone", "4001234567"),
    ]


def test_contact_extraction_normalizes_grouped_mobile_number():
    from core.mobile.collect.contacts import extract_contacts

    contacts = extract_contacts("有意向者联系邵部长 136 9362 0306。")

    assert [(item["channel"], item["value"]) for item in contacts] == [
        ("phone", "13693620306")
    ]
    assert "邵部长" in contacts[0]["context"]


def test_contact_extraction_does_not_treat_numeric_email_as_mobile():
    from core.mobile.collect.contacts import extract_contacts

    contacts = extract_contacts("简历投递邮箱：15904896909@163.com")

    assert [(item["channel"], item["value"]) for item in contacts] == [
        ("email", "15904896909@163.com")
    ]


def test_contact_extraction_normalizes_parenthesized_area_code():
    from core.mobile.collect.contacts import extract_contacts

    contacts = extract_contacts(
        "公示时间为5月8日至13日，受理电话：（010）63072558，"
        "来信地址为新华社人事局。"
    )

    assert [(item["channel"], item["value"]) for item in contacts] == [
        ("telephone", "010-63072558")
    ]
    assert "新华社人事局" in contacts[0]["context"]


def test_visual_contact_validation_rejects_domain_mislabeled_as_email():
    from core.mobile.collect.contacts import normalize_contact_candidate

    assert (
        normalize_contact_candidate(
            {"channel": "email", "value": "www.example.org", "context": "官网"}
        )
        is None
    )


def test_visual_contact_validation_accepts_anchored_local_telephone():
    from core.mobile.collect.contacts import normalize_contact_candidate

    contact = normalize_contact_candidate(
        {
            "channel": "telephone",
            "value": "58158252",
            "context": "项目咨询电话：58158252",
        }
    )

    assert contact is not None
    assert contact["value"] == "58158252"


def test_contextual_analysis_is_separate_from_source_version_fields():
    from api.models.mobile_collect import ExtractField
    from api.services.source_documents.service import (
        _analysis_fingerprint,
        _compact_contextual_analysis,
        _source_analysis,
    )

    capture = _capture(raw_html=b"raw", rendered_html=b"dom")
    source_analysis = _source_analysis(
        capture,
        contacts=[],
        image_analysis=[],
    )
    assert source_analysis["scope"] == "source"
    assert "target_id" not in source_analysis
    assert "content" not in source_analysis["fields"]

    compact = _compact_contextual_analysis(
        {
            "fields": {
                "title": "文章标题",
                "content": "完整正文",
                "bid_deadline": "2026-08-01",
            },
            "score": 90,
        }
    )
    assert compact["fields"] == {"bid_deadline": "2026-08-01"}

    fields = [ExtractField(name="bid_deadline", description="截止时间")]
    first = _analysis_fingerprint(
        version_id="version-1",
        target_id="target-1",
        target_name="公司甲",
        keyword="招标",
        fields=fields,
    )
    second = _analysis_fingerprint(
        version_id="version-1",
        target_id="target-2",
        target_name="公司乙",
        keyword="招标",
        fields=fields,
    )
    assert first != second


def test_contextual_analysis_fingerprint_tracks_prompt_content(monkeypatch):
    from api.models.mobile_collect import ExtractField
    from api.services.source_documents import service

    fields = [ExtractField(name="summary", description="摘要")]
    monkeypatch.setattr(
        service,
        "article_analysis_prompt_fingerprint",
        lambda: "prompt-version-one",
    )
    first = service._analysis_fingerprint(
        version_id="version-1",
        target_id="target-1",
        target_name="公司甲",
        keyword="招标",
        fields=fields,
    )
    monkeypatch.setattr(
        service,
        "article_analysis_prompt_fingerprint",
        lambda: "prompt-version-two",
    )
    second = service._analysis_fingerprint(
        version_id="version-1",
        target_id="target-1",
        target_name="公司甲",
        keyword="招标",
        fields=fields,
    )

    assert first != second


def test_finding_context_identity_is_stable_and_finding_scoped():
    from api.dao.finding_contexts import context_id_for_finding

    assert context_id_for_finding("finding-1") == context_id_for_finding("finding-1")
    assert context_id_for_finding("finding-1") != context_id_for_finding("finding-2")


def test_finding_context_agent_normalizes_singleton_provider_wrapper():
    from api.services.finding_context.agent import parse_finding_context_result

    parsed = parse_finding_context_result([{"title": "测试上下文"}])

    assert parsed.title == "测试上下文"


def test_finding_context_sanitizes_untrusted_evidence_references():
    from api.services.finding_context.schemas import (
        ContextFact,
        ContextNarrative,
        ContextVisualFinding,
        FindingContextResult,
    )
    from api.services.finding_context.service import sanitize_agent_result

    result = FindingContextResult(
        title="测试上下文",
        overview=ContextNarrative(
            text="证据化概览",
            kind="fact",
            confidence=90,
            evidence_refs=["finding:f-1", "source:invented"],
        ),
        business_background=ContextNarrative(
            text="缺少证据的背景",
            kind="fact",
            confidence=95,
        ),
        key_facts=[
            ContextFact(
                statement="可确认事实",
                confidence=100,
                evidence_refs=["finding:f-1", "image:invented"],
            ),
            ContextFact(
                statement="无证据内容",
                confidence=100,
                evidence_refs=["source:invented"],
            ),
        ],
        visual_findings=[
            ContextVisualFinding(
                evidence_ref="image:valid",
                summary="有效视觉证据",
            ),
            ContextVisualFinding(
                evidence_ref="image:invented",
                summary="无效视觉证据",
            ),
        ],
    )

    payload = sanitize_agent_result(
        result,
        allowed_refs={"finding:f-1", "image:valid"},
        fallback_title="",
        fallback_overview="",
    )

    assert payload["overview"]["evidence_refs"] == ["finding:f-1"]
    assert payload["business_background"]["kind"] == "inference"
    assert payload["business_background"]["confidence"] <= 40
    assert len(payload["key_facts"]) == 1
    assert payload["key_facts"][0]["evidence_refs"] == ["finding:f-1"]
    assert [item["evidence_ref"] for item in payload["visual_findings"]] == [
        "image:valid"
    ]


def test_finding_context_prompt_is_runtime_critical():
    from api.services.library_runtime import CORE_PROMPT_SLUGS

    assert "finding_context/organizer" in CORE_PROMPT_SLUGS
    assert "source_document/contact_attribution" in CORE_PROMPT_SLUGS


def test_finding_context_worker_wakes_for_incremental_enqueue(monkeypatch):
    import asyncio
    from contextlib import suppress

    from api.services.finding_context import service

    async def scenario():
        pending = []
        processed = asyncio.Event()

        async def claim_next_pending(_db):
            return pending.pop(0) if pending else None

        async def process_claimed(_db, context):
            assert context["finding_id"] == "finding-late"
            processed.set()

        monkeypatch.setattr(
            service.context_dao,
            "claim_next_pending",
            claim_next_pending,
        )
        monkeypatch.setattr(service, "_process_claimed", process_claimed)
        monkeypatch.setattr(service, "_worker_wakeup", asyncio.Event())

        worker = asyncio.create_task(service._worker_loop(object()))
        await asyncio.sleep(0)
        await asyncio.sleep(0)
        assert not worker.done()

        pending.append({"finding_id": "finding-late"})
        service._worker_wakeup.set()
        await asyncio.wait_for(processed.wait(), timeout=1)
        assert not worker.done()

        worker.cancel()
        with suppress(asyncio.CancelledError):
            await worker

    asyncio.run(scenario())


def test_source_document_prompts_reject_multi_entity_roundups():
    from Sere1nGraph.graph.prompts.loader import load_prompt
    from api.services.source_documents.analysis import (
        ARTICLE_ANALYSIS_PROMPT_SLUG,
        CONTACT_ATTRIBUTION_PROMPT_SLUG,
        RELEVANCE_REVIEW_PROMPT_SLUG,
    )

    extraction_prompt = load_prompt(ARTICLE_ANALYSIS_PROMPT_SLUG)
    review_prompt = load_prompt(RELEVANCE_REVIEW_PROMPT_SLUG)
    contact_prompt = load_prompt(CONTACT_ATTRIBUTION_PROMPT_SLUG)

    assert "multi_entity_roundup" in extraction_prompt
    assert "不得超过 39" in extraction_prompt
    assert "目标只是其中一个条目" in review_prompt
    assert "此类必须 `reject`" in review_prompt
    assert "独立相关性审核员" in review_prompt
    assert "不得依据常识" in contact_prompt
    assert "第三方公众号" in contact_prompt
    assert "必须逐项返回本批全部候选" in contact_prompt


def test_contact_attribution_batches_candidates_with_bounded_concurrency(
    monkeypatch,
):
    import asyncio
    import json
    from contextlib import nullcontext
    from types import SimpleNamespace

    from api.services.source_documents import analysis

    active = 0
    max_active = 0
    batches: list[list[str]] = []

    class Structured:
        async def ainvoke(self, messages):
            nonlocal active, max_active
            payload = messages[-1].content.split("候选联系方式：\n", 1)[1]
            candidates = json.loads(payload)
            batches.append([item["candidate_id"] for item in candidates])
            active += 1
            max_active = max(max_active, active)
            await asyncio.sleep(0.01)
            active -= 1
            return analysis.ContactAttributionBatch(
                items=[
                    analysis.ContactAttributionDecision(
                        candidate_id=item["candidate_id"],
                        belongs_to_target=item["candidate_id"] != "contact_4",
                        confidence=95,
                        reason="正文明确归属目标",
                    )
                    for item in candidates
                ]
            )

    class Llm:
        def with_structured_output(self, _schema):
            return Structured()

    async def _config():
        return SimpleNamespace(
            runtime=SimpleNamespace(models=SimpleNamespace(default="test-model"))
        )

    monkeypatch.setattr(analysis, "_CONTACT_ATTRIBUTION_BATCH_SIZE", 2)
    monkeypatch.setattr(analysis, "_CONTACT_ATTRIBUTION_CONCURRENCY", 2)
    monkeypatch.setattr(analysis, "get_runtime_app_config", _config)
    monkeypatch.setattr(analysis, "create_llm", lambda *_args, **_kwargs: Llm())
    monkeypatch.setattr(analysis, "load_prompt", lambda _slug: "{{target_name}}")
    monkeypatch.setattr(
        analysis,
        "observation_context",
        lambda **_kwargs: nullcontext(),
    )
    contacts = [
        {
            "channel": "telephone",
            "value": f"010-1234567{index}",
            "label": f"座机: 010-1234567{index}",
            "context": "目标单位联系人",
            "source": "text",
        }
        for index in range(5)
    ]

    selected, error = asyncio.run(
        analysis.attribute_target_contacts(
            target_name="目标单位",
            target_aliases=[],
            title="联系方式公告",
            account="目标单位",
            summary="目标单位联系人名单",
            contacts=contacts,
            image_analysis=[],
        )
    )

    assert batches == [
        ["contact_0", "contact_1"],
        ["contact_2", "contact_3"],
        ["contact_4"],
    ]
    assert max_active == 2
    assert [item["value"] for item in selected] == [
        item["value"] for item in contacts[:4]
    ]
    assert error == ""


def test_contact_attribution_retries_incomplete_batch(monkeypatch):
    import asyncio
    import json
    from contextlib import nullcontext
    from types import SimpleNamespace

    from api.services.source_documents import analysis

    calls = 0

    class Structured:
        async def ainvoke(self, messages):
            nonlocal calls
            calls += 1
            payload = messages[-1].content.split("候选联系方式：\n", 1)[1]
            candidates = json.loads(payload)
            if calls == 1:
                candidates = candidates[:1]
            return analysis.ContactAttributionBatch(
                items=[
                    analysis.ContactAttributionDecision(
                        candidate_id=item["candidate_id"],
                        belongs_to_target=True,
                        confidence=95,
                        reason="正文明确归属目标",
                    )
                    for item in candidates
                ]
            )

    class Llm:
        def with_structured_output(self, _schema):
            return Structured()

    async def _config():
        return SimpleNamespace(
            runtime=SimpleNamespace(models=SimpleNamespace(default="test-model"))
        )

    monkeypatch.setattr(analysis, "_CONTACT_ATTRIBUTION_BATCH_SIZE", 24)
    monkeypatch.setattr(analysis, "_CONTACT_ATTRIBUTION_CONCURRENCY", 1)
    monkeypatch.setattr(analysis, "_CONTACT_ATTRIBUTION_ATTEMPTS", 2)
    monkeypatch.setattr(analysis, "get_runtime_app_config", _config)
    monkeypatch.setattr(analysis, "create_llm", lambda *_args, **_kwargs: Llm())
    monkeypatch.setattr(analysis, "load_prompt", lambda _slug: "{{target_name}}")
    monkeypatch.setattr(
        analysis,
        "observation_context",
        lambda **_kwargs: nullcontext(),
    )
    contacts = [
        {
            "channel": "email",
            "value": f"person{index}@example.gov.cn",
            "label": f"邮箱: person{index}@example.gov.cn",
            "context": "目标单位联系人",
            "source": "text",
        }
        for index in range(2)
    ]

    selected, error = asyncio.run(
        analysis.attribute_target_contacts(
            target_name="目标单位",
            target_aliases=[],
            title="联系方式公告",
            account="目标单位",
            summary="目标单位联系人名单",
            contacts=contacts,
            image_analysis=[],
        )
    )

    assert calls == 2
    assert [item["value"] for item in selected] == [
        item["value"] for item in contacts
    ]
    assert error == ""


def test_contact_attribution_splits_a_persistently_incomplete_batch(monkeypatch):
    import asyncio
    import json
    from contextlib import nullcontext
    from types import SimpleNamespace

    from api.services.source_documents import analysis

    reviewed_batches: list[list[str]] = []

    class Structured:
        async def ainvoke(self, messages):
            payload = messages[-1].content.split("候选联系方式：\n", 1)[1]
            candidates = json.loads(payload)
            reviewed_batches.append([item["candidate_id"] for item in candidates])
            returned = candidates[:-1] if len(candidates) > 1 else candidates
            return analysis.ContactAttributionBatch(
                items=[
                    analysis.ContactAttributionDecision(
                        candidate_id=item["candidate_id"],
                        belongs_to_target=True,
                        confidence=95,
                        reason="正文明确归属目标",
                    )
                    for item in returned
                ]
            )

    class Llm:
        def with_structured_output(self, _schema):
            return Structured()

    async def _config():
        return SimpleNamespace(
            runtime=SimpleNamespace(models=SimpleNamespace(default="test-model"))
        )

    async def _sleep(_seconds):
        return None

    monkeypatch.setattr(analysis, "_CONTACT_ATTRIBUTION_BATCH_SIZE", 24)
    monkeypatch.setattr(analysis, "_CONTACT_ATTRIBUTION_CONCURRENCY", 2)
    monkeypatch.setattr(analysis, "_CONTACT_ATTRIBUTION_ATTEMPTS", 2)
    monkeypatch.setattr(analysis, "get_runtime_app_config", _config)
    monkeypatch.setattr(analysis, "create_llm", lambda *_args, **_kwargs: Llm())
    monkeypatch.setattr(analysis, "load_prompt", lambda _slug: "{{target_name}}")
    monkeypatch.setattr(analysis, "observation_context", lambda **_kwargs: nullcontext())
    monkeypatch.setattr(analysis.asyncio, "sleep", _sleep)
    contacts = [
        {
            "channel": "email",
            "value": f"person{index}@example.gov.cn",
            "context": "目标单位联系人",
            "source": "text",
        }
        for index in range(4)
    ]

    selected, error = asyncio.run(
        analysis.attribute_target_contacts(
            target_name="目标单位",
            target_aliases=[],
            title="联系方式公告",
            account="目标单位",
            summary="目标单位联系人名单",
            contacts=contacts,
            image_analysis=[],
        )
    )

    assert len(selected) == 4
    assert error == ""
    assert ["contact_0"] in reviewed_batches
    assert ["contact_3"] in reviewed_batches


def test_contact_attribution_uses_single_object_schema_for_one_candidate(
    monkeypatch,
):
    import asyncio
    from contextlib import nullcontext
    from types import SimpleNamespace

    from api.services.source_documents import analysis

    invoked: list[type] = []

    class Structured:
        def __init__(self, schema):
            self.schema = schema

        async def ainvoke(self, _messages):
            invoked.append(self.schema)
            if self.schema is analysis.ContactAttributionBatch:
                return analysis.ContactAttributionBatch(items=[])
            return analysis.ContactAttributionDecision(
                candidate_id="contact_0",
                belongs_to_target=True,
                confidence=96,
                reason="正文明确归属目标",
            )

    class Llm:
        def with_structured_output(self, schema):
            return Structured(schema)

    async def _config():
        return SimpleNamespace(
            runtime=SimpleNamespace(models=SimpleNamespace(default="test-model"))
        )

    async def _sleep(_seconds):
        return None

    monkeypatch.setattr(analysis, "_CONTACT_ATTRIBUTION_ATTEMPTS", 2)
    monkeypatch.setattr(analysis, "get_runtime_app_config", _config)
    monkeypatch.setattr(analysis, "create_llm", lambda *_args, **_kwargs: Llm())
    monkeypatch.setattr(analysis, "load_prompt", lambda _slug: "{{target_name}}")
    monkeypatch.setattr(analysis, "observation_context", lambda **_kwargs: nullcontext())
    monkeypatch.setattr(analysis.asyncio, "sleep", _sleep)

    selected, error = asyncio.run(
        analysis.attribute_target_contacts(
            target_name="目标单位",
            target_aliases=[],
            title="联系方式公告",
            account="目标单位",
            summary="目标单位联系人",
            contacts=[
                {
                    "channel": "email",
                    "value": "person@example.gov.cn",
                    "context": "目标单位联系人",
                    "source": "text",
                }
            ],
            image_analysis=[],
        )
    )

    assert invoked == [
        analysis.ContactAttributionBatch,
        analysis.ContactAttributionDecision,
    ]
    assert len(selected) == 1
    assert error == ""


def test_article_scope_caps_prevent_single_roundup_item_from_passing():
    from api.services.source_documents.analysis import apply_article_scope_cap

    assert apply_article_scope_cap(98, "target_focused") == 98
    assert apply_article_scope_cap(98, "multi_entity_roundup") == 39
    assert apply_article_scope_cap(98, "incidental") == 20
    assert apply_article_scope_cap(98, "unknown") == 69


def test_relevance_review_requires_contact_attribution_agreement():
    from api.services.source_documents.analysis import apply_relevance_review

    capture = _capture(raw_html=b"raw", rendered_html=b"dom")
    draft = {
        "fields": {"summary": "目标项目联系人"},
        "score": 95,
        "subject_match": 96,
        "article_scope": "target_focused",
        "target_contact_values": ["13800138000"],
    }
    review = {
        "decision": "accept",
        "article_scope": "target_focused",
        "score": 92,
        "subject_match": 94,
        "summary": "文章主要介绍目标项目。",
        "target_contact_values": [],
        "reason": "正文聚焦目标，但未确认联系方式归属。",
    }

    result = apply_relevance_review(capture, draft, review)

    assert result["review_decision"] == "accept"
    assert result["score"] == 92
    assert result["subject_match"] == 94
    assert result["target_contacts"] == []
    assert result["fields"]["contact"] == ""


def test_target_contact_filter_can_restore_dual_agent_declared_wechat_name():
    from core.mobile.collect.contacts import extract_contacts
    from api.services.source_documents.analysis import filter_target_contacts

    text = "申请方式：微信：添加官方微信号 AHTV文体中心，备注拍客申请。"

    contacts = filter_target_contacts(
        extract_contacts(text),
        ["AHTV文体中心"],
        text=text,
    )

    assert contacts == [
        {
            "channel": "wechat",
            "value": "AHTV文体中心",
            "label": "微信号: AHTV文体中心",
            "context": text,
            "contexts": [text],
            "source": "text",
            "attribution": "dual_agent_declared",
        }
    ]


def test_target_contact_filter_accepts_labeled_model_values():
    from core.mobile.collect.contacts import extract_contacts
    from api.services.source_documents.analysis import filter_target_contacts

    text = "联系人：陈老师，联系电话：0756-2191815，邮箱：hr＠example.com。"
    contacts = filter_target_contacts(
        extract_contacts(text),
        ["联系电话：0756-2191815", "E-mail：hr＠example.com"],
        text=text,
    )

    assert [(item["channel"], item["value"]) for item in contacts] == [
        ("email", "hr@example.com"),
        ("telephone", "0756-2191815"),
    ]


def test_contextual_completion_uses_contact_attribution_for_image_evidence(
    monkeypatch,
):
    import asyncio

    from api.services.source_documents import service

    candidate = {
        "channel": "telephone",
        "value": "010-68377160",
        "label": "座机: 010-68377160",
        "context": "王老师：010-68377160",
        "source": "image",
        "image_index": 1,
    }

    async def attribute(**kwargs):
        assert kwargs["target_name"] == "目标研究院"
        return [{**candidate, "attribution": "contact_review_agent"}], ""

    monkeypatch.setattr(service, "attribute_target_contacts", attribute)
    result = asyncio.run(
        service._complete_contextual_analysis(
            {
                "fields": {"summary": "目标研究院招聘"},
                "relevance_review": {"target_contact_values": []},
            },
            capture=_capture(raw_html=b"raw", rendered_html=b"dom"),
            contacts=[candidate],
            image_analysis=[{"index": 1, "visible_text": "目标研究院招聘"}],
            target_name="目标研究院",
            target_aliases=[],
            project_id="project-1",
            task_id="task-1",
        )
    )

    assert result["target_contacts"][0]["value"] == "010-68377160"
    assert result["fields"]["contact"] == "座机: 010-68377160"


def test_contextual_completion_contact_auditor_is_authoritative(
    monkeypatch,
):
    import asyncio

    from api.services.source_documents import service

    contacts = [
        {
            "channel": "telephone",
            "value": "010-68377160",
            "label": "座机: 010-68377160",
            "context": "招聘电话：010-68377160",
            "source": "text",
        },
        {
            "channel": "email",
            "value": "hr@example.com",
            "label": "邮箱: hr@example.com",
            "context": "招聘邮箱：hr@example.com",
            "source": "text",
        },
    ]

    async def attribute(**kwargs):
        assert [item["value"] for item in kwargs["contacts"]] == [
            "010-68377160",
            "hr@example.com",
        ]
        return [{**kwargs["contacts"][1], "attribution": "contact_review_agent"}], ""

    monkeypatch.setattr(service, "attribute_target_contacts", attribute)
    result = asyncio.run(
        service._complete_contextual_analysis(
            {
                "fields": {"summary": "目标招聘"},
                "relevance_review": {
                    "target_contact_values": ["联系电话：010-68377160"]
                },
            },
            capture=_capture(raw_html=b"raw", rendered_html=b"dom"),
            contacts=contacts,
            image_analysis=[],
            target_name="目标研究院",
            target_aliases=[],
            project_id="project-1",
            task_id="task-1",
        )
    )

    assert [item["value"] for item in result["target_contacts"]] == [
        "hr@example.com"
    ]


def test_contextual_completion_falls_back_to_review_on_auditor_error(monkeypatch):
    import asyncio

    from api.services.source_documents import service

    candidate = {
        "channel": "telephone",
        "value": "010-68377160",
        "label": "座机: 010-68377160",
        "context": "招聘电话：010-68377160",
        "source": "text",
    }

    async def attribute(**_kwargs):
        return [], "模型暂不可用"

    monkeypatch.setattr(service, "attribute_target_contacts", attribute)
    result = asyncio.run(
        service._complete_contextual_analysis(
            {
                "fields": {"summary": "目标招聘"},
                "relevance_review": {
                    "target_contact_values": ["联系电话：010-68377160"]
                },
            },
            capture=_capture(raw_html=b"raw", rendered_html=b"dom"),
            contacts=[candidate],
            image_analysis=[],
            target_name="目标研究院",
            target_aliases=[],
            project_id="project-1",
            task_id="task-1",
        )
    )

    assert result["target_contacts"][0]["value"] == "010-68377160"
    assert result["contact_attribution_error"] == "模型暂不可用"


def test_contextual_completion_keeps_successful_contact_batches(monkeypatch):
    import asyncio

    from api.services.source_documents import service

    reviewed = {
        "channel": "telephone",
        "value": "010-68377160",
        "label": "座机: 010-68377160",
        "context": "招聘电话：010-68377160",
        "source": "text",
    }
    attributed = {
        "channel": "email",
        "value": "hr@example.gov.cn",
        "label": "邮箱: hr@example.gov.cn",
        "context": "招聘邮箱：hr@example.gov.cn",
        "source": "text",
        "attribution": "contact_review_agent",
    }

    async def attribute(**_kwargs):
        return [attributed], "另一个审核批次超时"

    monkeypatch.setattr(service, "attribute_target_contacts", attribute)
    result = asyncio.run(
        service._complete_contextual_analysis(
            {
                "fields": {"summary": "目标招聘"},
                "relevance_review": {
                    "target_contact_values": ["联系电话：010-68377160"]
                },
            },
            capture=_capture(raw_html=b"raw", rendered_html=b"dom"),
            contacts=[reviewed, attributed],
            image_analysis=[],
            target_name="目标研究院",
            target_aliases=[],
            project_id="project-1",
            task_id="task-1",
        )
    )

    assert result["target_contacts"] == [attributed]
    assert result["contact_attribution_error"] == "另一个审核批次超时"


def test_relevance_review_conservatively_merges_two_agents():
    from api.services.source_documents.analysis import apply_relevance_review

    capture = _capture(raw_html=b"raw", rendered_html=b"dom")
    draft = {
        "fields": {"summary": "错误的整篇摘要"},
        "score": 95,
        "subject_match": 95,
        "article_scope": "target_focused",
        "target_contact_values": ["13800138000"],
    }
    review = {
        "decision": "reject",
        "article_scope": "multi_entity_roundup",
        "score": 70,
        "subject_match": 39,
        "summary": "目标仅出现在行业汇总的一个条目中。",
        "target_contact_values": [],
        "reason": "文章主体为多个单位的行业汇总。",
    }

    result = apply_relevance_review(capture, draft, review)

    assert result["review_decision"] == "reject"
    assert result["article_scope"] == "multi_entity_roundup"
    assert result["score"] == 70
    assert result["subject_match"] == 39
    assert result["fields"]["summary"] == "目标仅出现在行业汇总的一个条目中。"
    assert result["target_contacts"] == []


def test_article_analysis_runs_independent_relevance_reviewer(monkeypatch):
    import asyncio

    from api.services.source_documents import analysis

    calls: list[str] = []

    async def extract(*_args, **_kwargs):
        calls.append("extract")
        return {
            "fields": {"summary": "初稿", "content": "正文"},
            "score": 96,
            "subject_match": 95,
            "article_scope": "target_focused",
            "target_contact_values": ["13800138000"],
        }

    async def review(*_args, draft_analysis, **_kwargs):
        assert draft_analysis["fields"]["summary"] == "初稿"
        calls.append("review")
        return {
            "decision": "accept",
            "article_scope": "target_focused",
            "score": 92,
            "subject_match": 93,
            "summary": "审核后的目标专属摘要",
            "target_contact_values": ["13800138000"],
            "reason": "全文聚焦目标项目。",
        }

    monkeypatch.setattr(analysis, "analyze_article_fields", extract)
    monkeypatch.setattr(analysis, "review_article_relevance", review)

    result = asyncio.run(
        analysis.analyze_and_review_article(
            _capture(raw_html=b"raw", rendered_html=b"dom"),
            fields=[],
            target_name="目标单位",
            keyword="目标单位 招标",
            required_subject_match=70,
        )
    )

    assert calls == ["extract", "review"]
    assert result["review_decision"] == "accept"
    assert result["subject_match"] == 93
    assert result["fields"]["summary"] == "审核后的目标专属摘要"
    assert result["target_contacts"][0]["value"] == "13800138000"


def test_target_review_gate_requires_accept_decision_and_threshold():
    from api.services.source_documents.service import _passes_target_review

    assert _passes_target_review(
        {"review_decision": "accept", "subject_match": 90}, 70
    )
    assert not _passes_target_review(
        {"review_decision": "reject", "subject_match": 90}, 70
    )
    assert not _passes_target_review(
        {"review_decision": "accept", "subject_match": 69}, 70
    )
    assert not _passes_target_review({"subject_match": 100}, 70)


def test_contextual_link_compaction_keeps_final_contact_context_and_policy():
    from api.services.source_documents import service

    compact = service._compact_contextual_analysis(
        {
            "fields": {
                "summary": "目标单位招聘",
                "content": "完整正文不应复制到关联层",
                "contact": "座机: 010-63072558",
            },
            "target_contacts": [
                {
                    "channel": "telephone",
                    "value": "010-63072558",
                    "context": "受理电话，新华社人事局",
                }
            ],
        }
    )

    assert compact["fields"] == {"summary": "目标单位招聘"}
    assert compact["target_contact_values"] == ["010-63072558"]
    assert compact["target_contacts"][0]["context"] == "受理电话，新华社人事局"
    assert compact["contact_policy_version"] == service._CONTACT_POLICY_VERSION


def test_rejected_relevance_review_stops_before_source_persistence(monkeypatch):
    import asyncio

    from api.services.source_documents import service

    capture = _capture(raw_html=b"raw", rendered_html=b"dom")
    capture.images = []

    class _Provider:
        async def capture(self, *_args, **_kwargs):
            return capture

    async def get_version(*_args, **_kwargs):
        return None

    async def get_document_link(*_args, **_kwargs):
        return None

    async def reject(*_args, **_kwargs):
        return {
            "fields": {"summary": "目标仅为汇总中的一个条目"},
            "score": 60,
            "subject_match": 39,
            "score_reason": "文章主体为多个单位的行业汇总。",
            "article_scope": "multi_entity_roundup",
            "review_decision": "reject",
            "target_contact_values": [],
            "target_contacts": [],
        }

    async def forbidden(*_args, **_kwargs):
        raise AssertionError("审核拒绝后不得写入来源版本或 Target 关联")

    monkeypatch.setattr(service, "get_source_document_provider", lambda _url: _Provider())
    monkeypatch.setattr(service.source_dao, "get_version", get_version)
    monkeypatch.setattr(
        service.source_dao,
        "get_document_link",
        get_document_link,
    )
    monkeypatch.setattr(service, "analyze_and_review_article", reject)
    monkeypatch.setattr(service.source_dao, "begin_version", forbidden)
    monkeypatch.setattr(service.source_dao, "upsert_document", forbidden)
    monkeypatch.setattr(service.source_dao, "link_document", forbidden)

    result = asyncio.run(
        service.ingest_source_url(
            object(),
            url=capture.canonical_url,
            project_id="project-1",
            target={"target_id": "target-1", "canonical_name": "目标单位"},
            run_task_id="run-1",
            keyword="目标单位 招标",
            min_subject_match=70,
        )
    )

    assert result["ok"] is False
    assert result["rejected"] is True
    assert result["review_decision"] == "reject"
    assert result["article_scope"] == "multi_entity_roundup"


@pytest.mark.parametrize(
    "failure",
    [
        {"analysis_error": "结构化模型超时"},
        {"relevance_review": {"review_error": "审核模型限流"}},
    ],
)
def test_analysis_failure_is_retryable_instead_of_semantic_rejection(
    monkeypatch,
    failure,
):
    import asyncio

    from api.services.source_documents import service
    from api.services.source_documents.contracts import SourceDocumentAnalysisError

    capture = _capture(raw_html=b"raw", rendered_html=b"dom")

    class _Provider:
        async def capture(self, *_args, **_kwargs):
            return capture

    async def get_version(*_args, **_kwargs):
        return None

    async def failed_analysis(*_args, **_kwargs):
        return {
            "fields": {"summary": ""},
            "score": 0,
            "subject_match": 0,
            "review_decision": "reject",
            **failure,
        }

    async def forbidden(*_args, **_kwargs):
        raise AssertionError("技术失败不得写入来源版本或相关性拒绝关联")

    monkeypatch.setattr(service, "get_source_document_provider", lambda _url: _Provider())
    monkeypatch.setattr(service.source_dao, "get_version", get_version)
    monkeypatch.setattr(service, "analyze_and_review_article", failed_analysis)
    monkeypatch.setattr(service.source_dao, "begin_version", forbidden)
    monkeypatch.setattr(service.source_dao, "upsert_document", forbidden)
    monkeypatch.setattr(service.source_dao, "link_document", forbidden)

    with pytest.raises(SourceDocumentAnalysisError):
        asyncio.run(
            service.ingest_source_url(
                object(),
                url=capture.canonical_url,
                project_id="project-1",
                target={
                    "target_id": "target-1",
                    "canonical_name": "目标单位",
                },
                run_task_id="run-1",
                keyword="目标单位 招聘",
                min_subject_match=70,
            )
        )


def test_image_evidence_rechecks_a_text_only_rejection(monkeypatch):
    import asyncio

    from api.services.source_documents import service

    capture = _capture(raw_html=b"raw", rendered_html=b"dom")
    calls: list[str] = []

    class _Provider:
        async def capture(self, *_args, **_kwargs):
            return capture

    async def analyze(document, **_kwargs):
        if "【文章图片 OCR 与视觉证据】" not in document.text:
            calls.append("text")
            return {
                "fields": {"summary": "网页正文缺少主体"},
                "score": 20,
                "subject_match": 10,
                "review_decision": "reject",
                "article_scope": "uncertain",
                "relevance_review": {"decision": "reject"},
            }
        calls.append("image")
        assert "中国航天空气动力技术研究院" in document.text
        return {
            "fields": {"summary": "招聘海报明确属于目标研究院"},
            "score": 96,
            "subject_match": 98,
            "review_decision": "accept",
            "article_scope": "target_focused",
            "relevance_review": {"decision": "accept"},
        }

    async def analyze_images(*_args, **_kwargs):
        return (
            [
                {
                    "index": 0,
                    "description": "目标研究院校园招聘海报",
                    "visible_text": "中国航天空气动力技术研究院 2026届校园招聘",
                    "contacts": [
                        {
                            "channel": "email",
                            "value": "hr@example.com",
                            "context": "简历投递 hr@example.com",
                        }
                    ],
                    "is_key_evidence": True,
                    "importance_score": 98,
                }
            ],
            "",
        )

    async def complete(analysis, **_kwargs):
        return analysis

    monkeypatch.setattr(service, "get_source_document_provider", lambda _url: _Provider())
    monkeypatch.setattr(service, "analyze_and_review_article", analyze)
    monkeypatch.setattr(service, "analyze_article_images", analyze_images)
    monkeypatch.setattr(service, "_complete_contextual_analysis", complete)

    result = asyncio.run(
        service.ingest_source_url(
            object(),
            url=capture.canonical_url,
            target={"target_id": "target-1", "canonical_name": "目标研究院"},
            extract_fields=[],
            persist=False,
            min_subject_match=70,
        )
    )

    assert calls == ["text", "image"]
    assert result["ok"] is True
    assert result["review_decision"] == "accept"
    assert result["image_evidence_used"] is True
    assert result["image_evidence_indices"] == [0]


def test_incomplete_capture_does_not_replace_better_existing_images():
    from dataclasses import replace

    from api.services.source_documents import service

    capture = replace(
        _capture(raw_html=b"raw", rendered_html=b"dom"),
        images=[],
        metadata={
            "image_urls": ["https://mmbiz.qpic.cn/evidence.jpg"],
            "image_download_errors": ["connection closed"],
        },
    )
    existing = {
        "media_policy_version": service._MEDIA_POLICY_VERSION - 1,
        "images": [
            {
                "source_url": "https://mmbiz.qpic.cn/evidence.jpg",
                "analysis": {"visible_text": "目标单位招聘"},
            }
        ],
        "capture_metadata": {
            "analyzed_image_urls": ["https://mmbiz.qpic.cn/evidence.jpg"]
        },
    }

    assert service._capture_has_more_complete_images(existing, capture) is False


def test_archive_completeness_treats_unsupported_svg_as_warning():
    from api.services.source_documents import service

    error, warning = service._split_image_analysis_diagnostics(
        "image index=0 content_type=image/svg+xml "
        "prepare_failed=UnidentifiedImageError"
    )
    status, messages = service._archive_completeness(
        capture_metadata={},
        image_analysis_error=error,
        image_analysis_warning=warning,
    )

    assert error == ""
    assert "image/svg+xml" in warning
    assert status == "complete_with_warnings"
    assert messages and "无效或不受支持" in messages[0]


def test_archive_completeness_treats_partial_screenshots_as_warning():
    from api.services.source_documents import service

    status, messages = service._archive_completeness(
        capture_metadata={
            "screenshot_capture_error": "截图超时，已保留前 3 张"
        },
        image_analysis_error="",
        image_analysis_warning="",
    )

    assert status == "complete_with_warnings"
    assert messages == ["页面截图部分完成: 截图超时，已保留前 3 张"]


def test_archive_completeness_marks_truncated_images_partial():
    from api.services.source_documents import service

    status, messages = service._archive_completeness(
        capture_metadata={"images_truncated": 3},
        image_analysis_error="",
        image_analysis_warning="",
    )

    assert status == "partial"
    assert messages == ["页面图片因安全上限截断 3 张"]


def test_archive_completeness_reports_missing_evidence_as_partial():
    from api.services.source_documents import service

    status, messages = service._archive_completeness(
        capture_metadata={
            "image_download_errors": ["image-1 timeout", "image-2 timeout"],
            "screenshot_capture_error": "Page.captureScreenshot timeout",
        },
        image_analysis_error="image index=2 analyze_failed=TimeoutError",
        image_analysis_warning="",
    )

    assert status == "partial"
    assert any("原图下载失败 2 张" in message for message in messages)
    assert any("页面截图不完整" in message for message in messages)
    assert any("图片识别失败" in message for message in messages)


def test_archive_completeness_keeps_permanent_remote_404_as_warning():
    from api.services.source_documents import service

    status, messages = service._archive_completeness(
        capture_metadata={
            "attachment_download_warnings": ["file.pdf: HTTP 404 Not Found"],
            "image_download_warnings": ["image.png: HTTP 410 Gone"],
        },
        image_analysis_error="",
        image_analysis_warning="",
    )

    assert status == "complete_with_warnings"
    assert messages == [
        "来源站点已有 1 张原图永久失效",
        "来源站点已有 1 个附件永久失效",
    ]


def test_capture_repairs_attachment_text_and_image_analysis():
    from api.services.source_documents import service
    from api.services.source_documents.contracts import CapturedAttachment

    capture = _capture(raw_html=b"raw", rendered_html=b"dom")
    capture.attachments = [
        CapturedAttachment(
            index=0,
            source_url="https://example.gov.cn/scan.pdf",
            filename="scan.pdf",
            data=b"pdf",
            content_type="application/pdf",
            extracted_text="OCR 后的完整正文",
            text_format="pdf",
        )
    ]
    version = {
        "media_policy_version": service._MEDIA_POLICY_VERSION,
        "capture_metadata": {
            "archived_attachment_urls": ["https://example.gov.cn/scan.pdf"]
        },
        "attachments": [
            {
                "source_url": "https://example.gov.cn/scan.pdf",
                "text_error": "PDF 未提取到可读文本，可能为扫描件",
            }
        ],
        "image_analysis_error": "模型结构化输出异常",
        "image_analysis": [{"index": 0, "description": "普通配图"}],
    }

    assert service._capture_has_more_complete_images(version, capture) is True


def test_ingest_uses_final_redirect_url_as_document_identity(monkeypatch):
    import asyncio

    from api.dao.source_documents import document_id_for_url
    from api.services.source_documents import service
    from api.services.source_documents.contracts import CapturedDocument

    capture = CapturedDocument(
        source_type="official_website_document",
        canonical_url="https://www.example.gov.cn/notices/final.html",
        requested_url="http://example.gov.cn/notices/legacy.html",
        title="公开通知",
        account="example.gov.cn",
        publish_time="2026-08-12",
        text="公开通知正文",
        raw_html=b"<html></html>",
        rendered_html=b"<html></html>",
    )

    class _Provider:
        async def capture(self, *_args, **_kwargs):
            return capture

    monkeypatch.setattr(
        service, "get_source_document_provider", lambda _url: _Provider()
    )
    result = asyncio.run(
        service.ingest_source_url(
            object(),
            url=capture.requested_url,
            persist=False,
            analysis_mode="trusted_official",
        )
    )

    assert result["source_url"] == capture.canonical_url
    assert result["document_id"] == document_id_for_url(capture.canonical_url)


def test_image_analysis_retries_transient_structured_output(monkeypatch):
    import asyncio
    from io import BytesIO
    from types import SimpleNamespace

    from PIL import Image

    from api.services.source_documents import analysis
    from api.services.source_documents.contracts import CapturedImage

    output = BytesIO()
    Image.new("RGB", (320, 180), "white").save(output, format="PNG")
    calls = 0

    class _Structured:
        async def ainvoke(self, _messages):
            nonlocal calls
            calls += 1
            if calls == 1:
                raise RuntimeError("invalid structured JSON")
            return analysis.ImageUnderstandingBatch(
                items=[
                    analysis.ImageUnderstanding(
                        index=0,
                        visible_text="联系电话 010-12345678",
                    )
                ]
            )

    class _Llm:
        def with_structured_output(self, _schema):
            return _Structured()

    async def _config():
        return SimpleNamespace(
            runtime=SimpleNamespace(models=SimpleNamespace(vision="v"))
        )

    async def _sleep(_seconds):
        return None

    monkeypatch.setattr(analysis, "get_runtime_app_config", _config)
    monkeypatch.setattr(
        analysis, "create_llm", lambda *_args, **_kwargs: _Llm()
    )
    monkeypatch.setattr(analysis.asyncio, "sleep", _sleep)
    items, errors = asyncio.run(
        analysis._analyze_image_batch(
            [
                CapturedImage(
                    index=0,
                    source_url="https://example.gov.cn/contact.png",
                    data=output.getvalue(),
                    content_type="image/png",
                )
            ],
            project_id="project-1",
            task_id="task-1",
        )
    )

    assert calls == 2
    assert errors == []
    assert items[0]["visible_text"] == "联系电话 010-12345678"


def test_image_analysis_accepts_top_level_array_fallback(monkeypatch):
    import asyncio
    from io import BytesIO
    from types import SimpleNamespace

    from PIL import Image

    from api.services.source_documents import analysis
    from api.services.source_documents.contracts import CapturedImage

    output = BytesIO()
    Image.new("RGB", (320, 180), "white").save(output, format="PNG")
    schemas: list[type] = []

    class _Structured:
        def __init__(self, schema):
            self.schema = schema

        async def ainvoke(self, _messages):
            if self.schema is analysis.ImageUnderstandingBatch:
                raise ValueError("Input should be an object; model returned list")
            return analysis.ImageUnderstandingList(
                [analysis.ImageUnderstanding(index=0, visible_text="目标证据")]
            )

    class _Llm:
        def with_structured_output(self, schema):
            schemas.append(schema)
            return _Structured(schema)

    async def _config():
        return SimpleNamespace(
            runtime=SimpleNamespace(models=SimpleNamespace(vision="v"))
        )

    async def _sleep(_seconds):
        return None

    monkeypatch.setattr(analysis, "get_runtime_app_config", _config)
    monkeypatch.setattr(analysis, "create_llm", lambda *_args, **_kwargs: _Llm())
    monkeypatch.setattr(analysis.asyncio, "sleep", _sleep)
    items, errors = asyncio.run(
        analysis._analyze_image_batch(
            [
                CapturedImage(
                    index=0,
                    source_url="https://example.gov.cn/evidence.png",
                    data=output.getvalue(),
                    content_type="image/png",
                )
            ],
            project_id="project-1",
            task_id="task-1",
            source_type="official_website_document",
        )
    )

    assert schemas == [
        analysis.ImageUnderstandingBatch,
        analysis.ImageUnderstandingList,
    ]
    assert errors == []
    assert items[0]["visible_text"] == "目标证据"


def test_rejected_review_refreshes_an_existing_discovery_link(monkeypatch):
    import asyncio

    from api.services.source_documents import service

    capture = _capture(raw_html=b"raw", rendered_html=b"dom")
    capture.images = []
    capture.text = "正文主体不是目标单位"
    linked: list[dict] = []

    class _Provider:
        async def capture(self, *_args, **_kwargs):
            return capture

    async def get_version(*_args, **_kwargs):
        return {
            "version_id": "version-existing",
            "status": "ready",
            "media_policy_version": service._MEDIA_POLICY_VERSION,
            "contact_policy_version": service._CONTACT_POLICY_VERSION,
            "archive_status": "complete",
            "capture_metadata": {"analyzed_image_urls": []},
            "contacts": [],
            "image_analysis": [],
        }

    async def get_document_link(*_args, **_kwargs):
        return {"document_id": "document-existing"}

    async def reject(*_args, **_kwargs):
        return {
            "fields": {"summary": "正文主体不是目标单位"},
            "score": 20,
            "subject_match": 15,
            "score_reason": "目标仅被顺带提及",
            "article_scope": "incidental_mention",
            "review_decision": "reject",
            "target_contact_values": [],
            "target_contacts": [],
        }

    async def complete(analysis, **_kwargs):
        return analysis

    async def link(*_args, **kwargs):
        linked.append(kwargs)

    monkeypatch.setattr(service, "get_source_document_provider", lambda _url: _Provider())
    monkeypatch.setattr(service.source_dao, "get_version", get_version)
    monkeypatch.setattr(
        service.source_dao,
        "get_document_link",
        get_document_link,
    )
    monkeypatch.setattr(service, "analyze_and_review_article", reject)
    monkeypatch.setattr(service, "_complete_contextual_analysis", complete)
    monkeypatch.setattr(service, "_link_discovery", link)

    result = asyncio.run(
        service.ingest_source_url(
            object(),
            url=capture.canonical_url,
            project_id="project-1",
            target={"target_id": "target-1", "canonical_name": "目标单位"},
            task_def_id="task-1",
            run_task_id="run-review-1",
            keyword="目标单位 联系方式",
            min_subject_match=70,
        )
    )

    assert result["rejected"] is True
    assert len(linked) == 1
    assert linked[0]["run_task_id"] == "run-review-1"
    assert linked[0]["score"] == 20
    assert linked[0]["subject_match"] == 15
    assert linked[0]["contextual_analysis"]["review_decision"] == "reject"


def test_source_detail_can_select_immutable_version(monkeypatch):
    import asyncio

    from api.services.source_documents import service

    async def _get_document(db, document_id):
        return {"document_id": document_id, "latest_version_id": "version-new"}

    async def _get_version(db, version_id):
        return {
            "document_id": "doc-1",
            "version_id": version_id,
            "status": "ready",
        }

    async def _get_latest_version(db, document_id):
        raise AssertionError("指定版本时不应读取 latest_version_id")

    async def _get_links(db, document_id, project_id=""):
        return [{"project_id": project_id, "document_id": document_id}]

    monkeypatch.setattr(service.source_dao, "get_document", _get_document)
    monkeypatch.setattr(service.source_dao, "get_version", _get_version)
    monkeypatch.setattr(
        service.source_dao,
        "get_latest_version",
        _get_latest_version,
    )
    monkeypatch.setattr(
        service.source_dao,
        "get_links_for_document",
        _get_links,
    )

    detail = asyncio.run(
        service.get_source_document_detail(
            object(),
            "doc-1",
            project_id="project-1",
            version_id="version-old",
        )
    )
    assert detail is not None
    assert detail["version"]["version_id"] == "version-old"


def test_source_analysis_normalizes_ratio_scores_to_percentage():
    from api.services.source_documents.analysis import clamp_score, normalize_scores

    assert normalize_scores(0.96, 1.0) == (96, 100)
    assert normalize_scores(82, 95) == (82, 95)
    assert normalize_scores("invalid", None) == (0, 0)
    assert clamp_score(101.2) == 100
    assert clamp_score("79.6") == 80
    assert clamp_score("invalid") == 0


def test_article_image_analysis_normalizes_model_importance(monkeypatch):
    import asyncio

    from api.services.source_documents import analysis
    from api.services.source_documents.contracts import CapturedImage

    async def _analyze(*args, **kwargs):
        return (
            [
                {
                    "index": 3,
                    "description": "招标联系人截图",
                    "visible_text": "联系电话 13800138000",
                    "contacts": [],
                    "is_key_evidence": True,
                    "importance_score": 88.7,
                    "archive_reason": "包含联系方式",
                }
            ],
            [],
        )

    monkeypatch.setattr(analysis, "_analyze_image_batch", _analyze)
    result, error = asyncio.run(
        analysis.analyze_article_images(
            [
                CapturedImage(
                    index=3,
                    source_url="https://example.com/evidence.jpg",
                    data=b"image",
                    content_type="image/jpeg",
                )
            ]
        )
    )

    assert error == ""
    assert result[0]["importance_score"] == 89
    assert result[0]["is_key_evidence"] is True


def test_article_image_analysis_uses_bounded_batch_concurrency(monkeypatch):
    import asyncio

    from api.services.source_documents import analysis
    from api.services.source_documents.contracts import CapturedImage

    active = 0
    max_active = 0

    async def _analyze(images, **_kwargs):
        nonlocal active, max_active
        active += 1
        max_active = max(max_active, active)
        await asyncio.sleep(0.01)
        active -= 1
        return ([{"index": images[0].index, "description": "证据"}], [])

    monkeypatch.setattr(analysis, "_IMAGE_ANALYSIS_CONCURRENCY", 4)
    monkeypatch.setattr(analysis, "_analyze_image_batch", _analyze)
    images = [
        CapturedImage(
            index=index,
            source_url=f"https://example.gov.cn/{index}.jpg",
            data=b"image",
            content_type="image/jpeg",
        )
        for index in range(10)
    ]

    result, error = asyncio.run(
        analysis.analyze_article_images(images, batch_size=1)
    )

    assert 1 < max_active <= 4
    assert len(result) == 10
    assert error == ""


def test_article_image_preflight_isolates_unsupported_svg():
    from io import BytesIO

    from PIL import Image

    from api.services.source_documents.analysis import _prepare_image_inputs
    from api.services.source_documents.contracts import CapturedImage

    output = BytesIO()
    Image.new("RGB", (320, 180), "white").save(output, format="PNG")
    prepared, errors = _prepare_image_inputs(
        [
            CapturedImage(
                index=0,
                source_url="https://example.com/icon.svg",
                data=b'<svg xmlns="http://www.w3.org/2000/svg"></svg>',
                content_type="image/svg+xml",
            ),
            CapturedImage(
                index=1,
                source_url="https://example.com/evidence.png",
                data=output.getvalue(),
                content_type="image/png",
            ),
        ]
    )

    assert [image.index for image, _media_type, _encoded in prepared] == [1]
    assert len(errors) == 1
    assert "index=0" in errors[0]


def test_article_image_preflight_skips_tracking_pixels():
    from io import BytesIO

    from PIL import Image

    from api.services.source_documents.analysis import _prepare_image_inputs
    from api.services.source_documents.contracts import CapturedImage

    output = BytesIO()
    Image.new("RGB", (1, 100), "white").save(output, format="PNG")
    prepared, errors = _prepare_image_inputs(
        [
            CapturedImage(
                index=0,
                source_url="https://example.com/tracking.png",
                data=output.getvalue(),
                content_type="image/png",
            )
        ]
    )

    assert prepared == []
    assert len(errors) == 1
    assert "prepare_failed=ImagePreflightSkip" in errors[0]


def test_visual_contact_schema_tolerates_unsupported_model_channel():
    from api.services.source_documents.analysis import (
        ImageUnderstanding,
        VisualContact,
    )

    item = ImageUnderstanding(
        index=0,
        contacts=[
            VisualContact(channel="bilibili", value="example", context="主页")
        ],
    )

    assert item.contacts[0].channel == "bilibili"


def test_wechat_screenshots_use_bounded_cdp_viewport_capture():
    import asyncio
    import base64

    from api.services.source_documents.wechat import WechatArticleProvider

    class _Session:
        def __init__(self):
            self.calls = 0
            self.detached = False

        async def send(self, method, params):
            assert method == "Page.captureScreenshot"
            assert params["captureBeyondViewport"] is False
            self.calls += 1
            return {"data": base64.b64encode(b"jpeg-frame").decode("ascii")}

        async def detach(self):
            self.detached = True

    class _Context:
        def __init__(self, session):
            self.session = session

        async def new_cdp_session(self, _page):
            return self.session

    class _Page:
        def __init__(self):
            self.session = _Session()
            self.context = _Context(self.session)
            self.y = 0

        async def evaluate(self, script):
            if "scrollTo" in script:
                self.y = 0
                return None
            if "const before" in script:
                before = self.y
                self.y = min(100, self.y + 82)
                return {"before": before, "maxY": 100}
            if script == "window.scrollY":
                return self.y
            raise AssertionError(script)

        async def wait_for_timeout(self, _milliseconds):
            return None

    page = _Page()
    screenshots, error = asyncio.run(
        WechatArticleProvider._capture_screenshots(page)
    )

    assert error == ""
    assert len(screenshots) == 3
    assert all(item.data == b"jpeg-frame" for item in screenshots)
    assert page.session.calls == 3
    assert page.session.detached is True


def test_source_document_lock_serializes_waiters_and_cleans_registry():
    import asyncio

    from api.services.source_documents import service

    active = 0
    max_active = 0

    async def worker():
        nonlocal active, max_active
        async with service._hold_document_lock("doc-lock-test"):
            active += 1
            max_active = max(max_active, active)
            await asyncio.sleep(0.005)
            active -= 1

    async def scenario():
        await asyncio.gather(*(worker() for _ in range(6)))

    asyncio.run(scenario())
    assert max_active == 1
    assert "doc-lock-test" not in service._document_locks
    assert "doc-lock-test" not in service._document_lock_users


def test_artifact_object_id_is_content_addressed_for_safe_retry():
    from api.services.source_documents.service import _artifact_object_id

    first = _artifact_object_id("version-1", "raw", b"dynamic-token-one")
    same = _artifact_object_id("version-1", "raw", b"dynamic-token-one")
    retry = _artifact_object_id("version-1", "raw", b"dynamic-token-two")

    assert first == same
    assert first != retry


def test_failed_evidence_upgrade_restores_previous_ready_version(monkeypatch):
    import asyncio

    from api.services.source_documents import service

    restored: list[dict] = []
    errors: list[str] = []

    async def mark_ready(_db, *, version_id, payload):
        restored.append({"version_id": version_id, **payload})

    async def mark_error(_db, version_id, error):
        errors.append(f"{version_id}:{error}")

    monkeypatch.setattr(service.source_dao, "mark_version_ready", mark_ready)
    monkeypatch.setattr(service.source_dao, "mark_version_error", mark_error)

    asyncio.run(
        service._record_version_failure(
            object(),
            version_id="version-1",
            error=RuntimeError("new image upload failed"),
            previous_ready_version={
                "_id": "mongo-id",
                "version_id": "version-1",
                "status": "ready",
                "error": "",
                "updated_at": "old-time",
                "artifacts": {"raw_html_object_id": "object-raw"},
            },
        )
    )

    assert errors == []
    assert restored == [
        {
            "version_id": "version-1",
            "artifacts": {"raw_html_object_id": "object-raw"},
        }
    ]


def test_failed_new_version_is_marked_error(monkeypatch):
    import asyncio

    from api.services.source_documents import service

    errors: list[str] = []

    async def mark_error(_db, version_id, error):
        errors.append(f"{version_id}:{error}")

    monkeypatch.setattr(service.source_dao, "mark_version_error", mark_error)

    asyncio.run(
        service._record_version_failure(
            object(),
            version_id="version-new",
            error=RuntimeError("capture persistence failed"),
            previous_ready_version=None,
        )
    )

    assert errors == ["version-new:capture persistence failed"]


def test_cached_contact_refresh_rewrites_only_structured_artifact(monkeypatch):
    import asyncio

    from api.services.source_documents import service

    structured_payloads: list[dict] = []
    version_payloads: list[dict] = []

    async def store_structured(structured, **_kwargs):
        structured_payloads.append(structured)
        return {
            "structured_object_id": "object-structured-new",
            "structured_url": "/objects/object-structured-new",
        }

    async def mark_ready(_db, *, version_id, payload):
        version_payloads.append(payload)
        return {"version_id": version_id, **payload}

    monkeypatch.setattr(service, "_store_structured_json", store_structured)
    monkeypatch.setattr(service.source_dao, "mark_version_ready", mark_ready)
    result = asyncio.run(
        service._refresh_cached_source_contacts(
            object(),
            capture=_capture(raw_html=b"raw", rendered_html=b"dom"),
            document_id="document-1",
            version_id="version-1",
            version={
                "content_hash": "content-hash",
                "source_type": "wechat_article",
                "identity": {"title": "标题"},
                "content": {"text": "正文"},
                "artifacts": {
                    "raw_html_object_id": "object-raw",
                    "structured_object_id": "object-structured-old",
                },
                "storage_object_ids": [
                    "object-raw",
                    "object-structured-old",
                ],
                "images": [],
                "screenshots": [],
                "image_analysis_error": (
                    "image index=0 content_type=image/svg+xml "
                    "prepare_failed=UnidentifiedImageError"
                ),
            },
            contacts=[
                {
                    "channel": "email",
                    "value": "hr@example.com",
                    "label": "邮箱: hr@example.com",
                    "context": "招聘邮箱 hr@example.com",
                }
            ],
            image_analysis=[],
            target_id="target-1",
            project_id="project-1",
        )
    )

    assert structured_payloads[0]["evidence"]["contacts"][0]["context"] == (
        "招聘邮箱 hr@example.com"
    )
    assert structured_payloads[0]["provenance"]["artifacts"] == {
        "raw_html_object_id": "object-raw"
    }
    assert version_payloads[0]["artifacts"]["raw_html_object_id"] == (
        "object-raw"
    )
    assert version_payloads[0]["artifacts"]["structured_object_id"] == (
        "object-structured-new"
    )
    assert result["storage_object_ids"] == [
        "object-raw",
        "object-structured-old",
        "object-structured-new",
    ]
    assert version_payloads[0]["image_analysis_error"] == ""
    assert "image/svg+xml" in version_payloads[0]["image_analysis_warning"]
    assert version_payloads[0]["archive_status"] == "complete_with_warnings"
    assert (
        structured_payloads[0]["evidence"]["media"]["archive_status"]
        == "complete_with_warnings"
    )


def test_image_archive_policy_keeps_only_contact_and_key_evidence():
    from api.services.source_documents.contracts import CapturedImage
    from api.services.source_documents.service import _select_archive_images

    images = [
        CapturedImage(
            index=index,
            source_url=f"https://example.com/{index}.jpg",
            data=b"x",
            content_type="image/jpeg",
        )
        for index in range(4)
    ]
    selected = _select_archive_images(
        images,
        [
            {"index": 0, "importance_score": 10, "visible_text": "装饰图片"},
            {"index": 1, "importance_score": 20, "visible_text": "电话 13800138000"},
            {"index": 2, "importance_score": 90, "is_key_evidence": True},
            {"index": 3, "importance_score": 69, "visible_text": "普通配图"},
        ],
    )

    assert [image.index for image in selected] == [1, 2]


def test_image_archive_policy_discards_low_value_key_and_tiny_decoration():
    from api.services.source_documents.contracts import CapturedImage
    from api.services.source_documents.service import _select_archive_images

    images = [
        CapturedImage(
            index=0,
            source_url="https://example.com/scene.jpg",
            data=b"scene",
            content_type="image/jpeg",
            width=1080,
            height=720,
        ),
        CapturedImage(
            index=1,
            source_url="https://example.com/icon.gif",
            data=b"icon",
            content_type="image/gif",
            width=94,
            height=94,
        ),
    ]

    selected = _select_archive_images(
        images,
        [
            {"index": 0, "is_key_evidence": True, "importance_score": 70},
            {"index": 1, "is_key_evidence": True, "importance_score": 95},
        ],
    )

    assert selected == []


def test_image_archive_policy_rejects_invalid_model_contact():
    from api.services.source_documents.contracts import CapturedImage
    from api.services.source_documents.service import _select_archive_images

    selected = _select_archive_images(
        [
            CapturedImage(
                index=0,
                source_url="https://example.com/footer.jpg",
                data=b"footer",
                content_type="image/jpeg",
                width=800,
                height=300,
            )
        ],
        [
            {
                "index": 0,
                "importance_score": 20,
                "contacts": [
                    {
                        "channel": "email",
                        "value": "www.example.org",
                        "context": "官方网站",
                    }
                ],
            }
        ],
    )

    assert selected == []


def test_contact_findings_batch_uses_one_ordered_false_bulk_write():
    import asyncio
    from types import SimpleNamespace

    from api.dao import findings as findings_dao
    from api.db.collections import FINDINGS_COLLECTION

    class _Collection:
        def __init__(self):
            self.operations = []
            self.ordered = None

        async def bulk_write(self, operations, *, ordered):
            self.operations = operations
            self.ordered = ordered
            return SimpleNamespace()

    class _DB:
        def __init__(self):
            self.collection = _Collection()

        def __getitem__(self, name):
            assert name == FINDINGS_COLLECTION
            return self.collection

    db = _DB()
    count = asyncio.run(
        findings_dao.upsert_contact_findings_batch(
            db,
            [
                {
                    "finding_id": "contact-1",
                    "project_id": "project-1",
                    "task_id": "task-1",
                    "target_id": "target-1",
                    "source_document_id": "document-1",
                    "evidence_ref": {"record_id": "website:document-1"},
                },
                {
                    "finding_id": "contact-2",
                    "project_id": "project-1",
                    "task_id": "task-1",
                    "target_id": "target-1",
                    "source_document_id": "document-1",
                    "evidence_ref": {"record_id": "website:document-1"},
                },
            ],
        )
    )

    assert count == 2
    assert db.collection.ordered is False
    assert len(db.collection.operations) == 2
    first = db.collection.operations[0]
    assert first._filter == {"finding_id": "contact-1"}
    assert first._doc["$addToSet"]["task_ids"] == "task-1"
    assert first._doc["$addToSet"]["evidence_refs"] == {
        "record_id": "website:document-1"
    }


def test_contact_finding_reconciliation_removes_only_stale_record_evidence():
    import asyncio
    from types import SimpleNamespace

    from api.dao import findings as findings_dao
    from api.db.collections import FINDINGS_COLLECTION

    class _Cursor:
        def __init__(self, values):
            self.values = values

        async def to_list(self, length=None):
            del length
            return [dict(value) for value in self.values]

    class _Collection:
        def __init__(self):
            self.docs = {
                "delete-me": {
                    "finding_id": "delete-me",
                    "project_id": "project-1",
                    "type": "contact",
                    "evidence_refs": [{"record_id": "record-1"}],
                    "latest_evidence_ref": {"record_id": "record-1"},
                },
                "keep-shared": {
                    "finding_id": "keep-shared",
                    "project_id": "project-1",
                    "type": "contact",
                    "evidence_refs": [
                        {"record_id": "record-2", "context": "旧证据"},
                        {"record_id": "record-1", "context": "待撤销"},
                    ],
                    "latest_evidence_ref": {
                        "record_id": "record-1",
                        "context": "待撤销",
                    },
                },
                "still-valid": {
                    "finding_id": "still-valid",
                    "project_id": "project-1",
                    "type": "contact",
                    "evidence_refs": [{"record_id": "record-1"}],
                    "latest_evidence_ref": {"record_id": "record-1"},
                },
                "legacy-shared": {
                    "finding_id": "legacy-shared",
                    "project_id": "project-1",
                    "type": "contact",
                    "evidence_refs": [{"record_id": "record-1"}],
                    "latest_evidence_ref": {"record_id": "record-1"},
                    "evidence": {"record_id": "record-2"},
                },
            }

        def find(self, *_args, **_kwargs):
            return _Cursor(self.docs.values())

        async def find_one(self, query, _projection=None):
            value = self.docs.get(query.get("finding_id"))
            return dict(value) if value else None

        async def update_one(self, query, update):
            value = self.docs.get(query.get("finding_id"))
            if not value:
                return SimpleNamespace(modified_count=0)
            before = repr(value)
            record_id = (
                (update.get("$pull") or {}).get("evidence_refs") or {}
            ).get("record_id")
            if record_id:
                value["evidence_refs"] = [
                    item
                    for item in value.get("evidence_refs") or []
                    if item.get("record_id") != record_id
                ]
            value.update(update.get("$set") or {})
            return SimpleNamespace(modified_count=int(before != repr(value)))

        async def delete_one(self, query):
            value = self.docs.get(query.get("finding_id"))
            if value and not value.get("evidence_refs"):
                self.docs.pop(query["finding_id"])
                return SimpleNamespace(deleted_count=1)
            return SimpleNamespace(deleted_count=0)

    class _DB:
        def __init__(self):
            self.collection = _Collection()

        def __getitem__(self, name):
            assert name == FINDINGS_COLLECTION
            return self.collection

    db = _DB()
    result = asyncio.run(
        findings_dao.reconcile_contact_findings_for_record(
            db,
            project_id="project-1",
            record_id="record-1",
            keep_finding_ids=["still-valid"],
        )
    )

    assert result == {"evidence_removed": 3, "findings_deleted": 1}
    assert "delete-me" not in db.collection.docs
    assert db.collection.docs["keep-shared"]["evidence_refs"] == [
        {"record_id": "record-2", "context": "旧证据"}
    ]
    assert db.collection.docs["keep-shared"]["latest_evidence_ref"] == {
        "record_id": "record-2",
        "context": "旧证据",
    }
    assert db.collection.docs["still-valid"]["evidence_refs"] == [
        {"record_id": "record-1"}
    ]
    assert "legacy-shared" in db.collection.docs
